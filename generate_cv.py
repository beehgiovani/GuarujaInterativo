from docx import Document
from docx.shared import Pt, Inches
from fpdf import FPDF
import os

def create_docx(filename):
    doc = Document()
    
    # Title
    title = doc.add_heading('Bruno Giovani Pereira', 0)
    
    # Contact Info
    p = doc.add_paragraph()
    p.add_run('Guarujá, SP | 25/05/1993\n').italic = True
    p.add_run('brunogp.corretor@gmail.com\n').italic = True
    
    # Summary
    doc.add_heading('Resumo Profissional', level=1)
    doc.add_paragraph(
        "Desenvolvedor Full-Stack e Consultor Imobiliário especialista em Inteligência Territorial. "
        "Autodidata em tecnologia, com sólida experiência na criação de ecossistemas digitais que integram geoprocessamento, "
        "automação de dados (Web Scraping/OSINT) e CRM. Foco no desenvolvimento de soluções voltadas para eficiência "
        "de mercado, arrecadação municipal e análise de viabilidade imobiliária."
    )
    
    # Hard Skills
    doc.add_heading('Principais Competências Técnicas (Hard Skills)', level=1)
    skills = [
        "Linguagens: JavaScript (ES6+), Python, HTML5, CSS3.",
        "Frameworks & Plataformas: Node.js, Electron, Capacitor, Firebase, Supabase.",
        "Geoprocessamento: Google Maps API (Advanced), Google Earth API, OpenStreetMap, StreetView API.",
        "Automação & Dados: Web Scraping (Selenium, BeautifulSoup, Requests), OSINT, Tratamento de dados sensíveis.",
        "IA: Integração com LLMs (Gemini AI), Chatbots inteligentes."
    ]
    for skill in skills:
        doc.add_paragraph(skill, style='List Bullet')
        
    # Projects
    doc.add_heading('Projetos de Destaque', level=1)
    
    doc.add_heading('Guarujá Geo Lab / GeoMap', level=2)
    doc.add_paragraph(
        "Plataforma de inteligência territorial com mapa interativo, camadas de zoneamento e CRM integrado. "
        "Digitalizou processos de análise, reduzindo tempo de consulta em 80%."
    )
    
    doc.add_heading('Sistema OSINT & Scraper (V11)', level=2)
    doc.add_paragraph(
        "Automação em Python para coleta de dados públicos e normalização de débitos e proprietários."
    )
    
    # Experience
    doc.add_heading('Experiência Profissional', level=1)
    doc.add_paragraph("Desenvolvedor & Fundador (Projeto Independente) | 2023 – Presente", style='List Bullet')
    doc.add_paragraph("Consultor Imobiliário (Broker) | Litoral de SP", style='List Bullet')
    
    # Education
    doc.add_heading('Formação Acadêmica', level=1)
    doc.add_paragraph("Superior em Farmácia | Cursando", style='List Bullet')
    doc.add_paragraph("Ensino Médio Completo", style='List Bullet')
    doc.add_paragraph("Desenvolvimento de Software (Autodidata)", style='List Bullet')
    
    doc.save(filename)

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 16)
        self.cell(0, 10, 'Bruno Giovani Pereira', 0, 1, 'C')
        self.set_font('Arial', '', 10)
        self.cell(0, 5, 'Guarujá, SP | brunogp.corretor@gmail.com', 0, 1, 'C')
        self.ln(10)

    def chapter_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.set_fill_color(240, 240, 240)
        self.cell(0, 8, title, 0, 1, 'L', 1)
        self.ln(4)

    def chapter_body(self, body, bullet=False):
        self.set_font('Arial', '', 11)
        if bullet:
            self.multi_cell(0, 7, f"- {body}")
        else:
            self.multi_cell(0, 7, body)
        self.ln(4)

def create_pdf(filename):
    pdf = PDF()
    pdf.add_page()
    
    pdf.chapter_title('Resumo Profissional')
    pdf.chapter_body(
        "Desenvolvedor Full-Stack e Consultor Imobiliário especialista em Inteligência Territorial. "
        "Autodidata em tecnologia, com sólida experiência na criação de ecossistemas digitais que integram geoprocessamento, "
        "automação de dados (Web Scraping/OSINT) e CRM."
    )
    
    pdf.chapter_title('Principais Competências Técnicas')
    skills = [
        "JavaScript (ES6+), Python, HTML5, CSS3.",
        "Node.js, Electron, Capacitor, Firebase, Supabase.",
        "Google Maps API, Google Earth, OpenStreetMap.",
        "Web Scraping (Selenium, BeautifulSoup), OSINT.",
        "IA: Integração com LLMs (Gemini AI)."
    ]
    for skill in skills:
        pdf.chapter_body(skill, bullet=True)
        
    pdf.chapter_title('Projetos de Destaque')
    pdf.chapter_body("Guarujá Geo Lab: Inteligência territorial e CRM integrado.")
    pdf.chapter_body("Scraper V11: Automação OSINT e normalização de dados.")
    
    pdf.chapter_title('Formação')
    pdf.chapter_body("Superior em Farmácia (Cursando)", bullet=True)
    pdf.chapter_body("Desenvolvedor de Software (Autodidata)", bullet=True)
    
    pdf.output(filename)

if __name__ == "__main__":
    create_docx("Curriculo_Bruno_Pereira.docx")
    create_pdf("Curriculo_Bruno_Pereira.pdf")
    print("Arquivos criados com sucesso!")
