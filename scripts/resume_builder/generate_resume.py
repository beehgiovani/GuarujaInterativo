from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# --- RESUME DATA ---
NAME = "Bruno Giovani Pereira"
LOCATION = "Guarujá, SP"
CONTACTS = "brunogp.corretor@gmail.com | (13) 99124-8146"
LINKEDIN = "https://www.linkedin.com/in/beehgiovani/"
GITHUB = "https://github.com/beehgiovani/"

SUMMARY = (
    "Software Engineer full-stack com especialização em Geotecnologia (GIS), Automação Avançada (OSINT/LegalTech) "
    "e Desenvolvimento Mobile Native. Expertise em arquiteturas de alto impacto: plataformas GIS de inteligência imobiliária, "
    "sistemas de extração de Big Data com modelos OCR customizados (PyTorch) e PWAs de escala industrial. "
    "Criador do GuaruGeo e de soluções para clientes de grande porte como a Cimed."
)

EXPERIENCES = [
    {
        "title": "Senior Consultant (Kotlin) | Soo Tech",
        "period": "jan. 2024 - Presente",
        "description": "Prestacao de servicos de alto valor em Kotlin para cliente Agilidade via Soo Tech.",
        "bullets": [
            "Otimizacao de modulos criticos de performance em Kotlin, reduzindo latencia de processamento.",
            "Implementacao de novas funcionalidades seguindo metodologias ageis para entregas rapidas e de qualidade.",
            "Aplicacao de padroes avancados de design em Kotlin para escalabilidade e manutenibilidade.",
            "Colaboracao com equipe de produto para traduzir requisitos de negocio em solucoes tecnicas."
        ]
    },
    {
        "title": "Founder & Lead Developer | GuaruGeo (Guarujá Geo Lab)",
        "link": "https://guarujainterativo.web.app/",
        "period": "2023 - Presente",
        "description": "Plataforma GIS de inteligência imobiliária para o litoral paulista.",
        "bullets": [
            "Arquitetura Full-Stack: React, Node.js, Supabase (PostgreSQL + Realtime) e Firebase Hosting.",
            "Visualização GeoJSON com Google Maps API e Overpass API; índice espacial com RBush para alta performance.",
            "Motor OSINT para identificação de titularidade, situação jurídica e enriquecimento de dados proprietários.",
            "Admin Dashboard com analytics em tempo real, CRM integrado e emissão de certidões jurídicas.",
            "Site ao vivo: https://guarujainterativo.web.app/"
        ]
    },
    {
        "title": "Lead Developer | Cimed Experience",
        "period": "2024 - 2025",
        "description": "PWA de marketing e engajamento de escala industrial para a Cimed, líder nacional em produtos farmacêuticos.",
        "bullets": [
            "Stack completo: React 19, TypeScript, Vite, Framer Motion, Zustand, Supabase e Firebase.",
            "Social Hub com rankings em tempo real, feed gamificado e missões de Drive-to-Store.",
            "IA 'Claud.ia' como assistente de saúde e produto; integração de video scraping e PWA sem loja de apps.",
            "Site ao vivo: https://cimedexperience.web.app/"
        ]
    },
    {
        "title": "Especialista em Automação & Data Engineering",
        "period": "2021 - Presente",
        "description": "Esteiras de extração de dados complexos e automação de processos críticos (LegalTech/GovTech).",
        "bullets": [
            "Scraper Imobiliário (2 estágios): HTTP + Playwright com Tor multi-porta e OCR customizado (PyTorch/CRNN) para CAPTCHAs.",
            "Certibot (LegalTech): Bot de emissão automática de certidões em 6+ portais (TRF3, TJSP, TST, TRT2, Receita Federal, CJF).",
            "GIS Data Mining: Extração e normalização de polígonos de zoneamento (PDDU) a partir de documentos PDF/CAD.",
            "Automação de Certidões (Web App): Sistema com Supabase Edge Functions e InfoSimples API para gestão jurídica."
        ]
    },
    {
        "title": "Desenvolvedor Mobile Native Android",
        "period": "2021 - 2023",
        "description": "Soluções mobile de alto padrão com Kotlin e integrações de IA/Hardware.",
        "bullets": [
            "NidusCare: MVVM + Hilt + Coroutines; Gemini Vision AI para análise de prescrições; Firebase backend.",
            "MarketPoster: ML para remoção de fundo (background removal); impressão via Bluetooth thermal.",
            "PharmaFlow: Ecossistema de delivery com App Nativo e Dashboard React/Firebase."
        ]
    }
]

SKILLS = {
    "Linguagens": "Javascript (React/Node), TypeScript, Python (Scraping/AI), Kotlin (Jetpack Compose), SQL.",
    "Cloud & Infra": "Supabase, Firebase, Docker, PostgreSQL, Linux Server, DigitalOcean.",
    "Frontend": "React 19, Vite, Framer Motion, Zustand, Leaflet.js, Google Maps API.",
    "Automação/AI": "PyTorch (CRNN OCR), Playwright, ddddocr, Tor, Gemini API, Vertex AI.",
    "Especialidades": "OSINT, GIS, LegalTech, PWA, Scraping Anti-Bot, Computer Vision, Gamification."
}

EDUCATION = [
    "Bacharelado em Farmácia (Em andamento) - Base para soluções HealthTech.",
    "Desenvolvimento de Sistemas - Especialização autônoma em Full-Stack e Engenharia de Dados."
]

def generate_docx(output_path):
    doc = Document()
    
    title = doc.add_heading(NAME, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{LOCATION}\n")
    p.add_run(f"{CONTACTS}\n")
    p.add_run(f"LinkedIn: {LINKEDIN}\n")
    p.add_run(f"GitHub: {GITHUB}")
    
    doc.add_heading('Resumo Profissional', level=1)
    doc.add_paragraph(SUMMARY)
    
    doc.add_heading('Experiência e Realizações Técnicas', level=1)
    for exp in EXPERIENCES:
        p = doc.add_paragraph()
        run = p.add_run(exp['title'])
        run.bold = True
        run.font.size = Pt(12)
        p.add_run(f" ({exp['period']})")
        doc.add_paragraph(exp['description'])
        for bullet in exp['bullets']:
            doc.add_paragraph(bullet, style='List Bullet')
            
    doc.add_heading('Expertise Técnica', level=1)
    for k, v in SKILLS.items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(v)
        
    doc.add_heading('Educação', level=1)
    for edu in EDUCATION:
        doc.add_paragraph(edu, style='List Bullet')
        
    doc.save(output_path)
    print(f"DOCX Generated: {output_path}")

output_dir = r"C:\Users\bruno\Downloads\Curriculo_Novo"
if not os.path.exists(output_dir): os.makedirs(output_dir)
generate_docx(os.path.join(output_dir, "Curriculo_Bruno_Pereira_2026.docx"))
