from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import os

# --- DATOS DO CURRÍCULO ---
NAME = "Bruno Giovani Pereira"
LOCATION = "Guarujá, SP"
CONTACTS = "brunogp.corretor@gmail.com | (13) 99124-8146"
LINKEDIN = "https://www.linkedin.com/in/beehgiovani/"
GITHUB = "https://github.com/beehgiovani/"

SUMMARY = (
    "Software Engineer Full-Stack com especialização em Geotecnologia (GIS), Automação Avançada (OSINT/LegalTech) "
    "e Desenvolvimento Mobile Native. Expertise em arquiteturas de alto impacto: plataformas GIS de inteligência imobiliária, "
    "sistemas de extração de Big Data com modelos OCR customizados (PyTorch) e PWAs de escala industrial. "
    "Certificações recentes em Angular, Power BI e Storytelling pela University of Chicago. Criador do GuaruGeo e Cimed Experience."
)

EXPERIENCES = [
    {
        "title": "Senior Consultant (Kotlin) | Soo Tech",
        "period": "Jan. 2024 - Presente",
        "description": "Prestação de serviços de alto valor em Kotlin para cliente Agilidade via Soo Tech.",
        "bullets": [
            "Otimização de módulos críticos de performance em Kotlin, reduzindo a latência de processamento em 30%.",
            "Implementação de novas funcionalidades seguindo metodologias ágeis (Scrum/Kanban).",
            "Aplicação de padrões avançados de design (SOLID) para escalabilidade e manutenibilidade.",
            "Colaboração direta com equipes de produto para traduzir requisitos de negócio em soluções técnicas robustas."
        ]
    },
    {
        "title": "Founder & Lead Developer | GuaruGeo (Guarujá Geo Lab)",
        "link": "https://guarujainterativo.web.app/",
        "period": "2023 - Presente",
        "description": "Plataforma GIS de inteligência imobiliária proprietária para o litoral paulista.",
        "bullets": [
            "Arquitetura Full-Stack: React 19, Node.js, Supabase (PostgreSQL + Realtime) e Firebase Hosting.",
            "Visualização GeoJSON avançada com Google Maps API e Overpass API; índice espacial RBush para performance extrema.",
            "Motor OSINT para identificação de titularidade, situação jurídica e enriquecimento de dados proprietários via API.",
            "Admin Dashboard com analytics em tempo real, CRM integrado e sistema automatizado de emissão de certidões.",
            "Link do Projeto: https://guarujainterativo.web.app/"
        ]
    },
    {
        "title": "Lead Developer | Cimed Experience",
        "link": "https://cimedexperience.web.app/",
        "period": "2024 - Março 2026",
        "description": "PWA de marketing e engajamento de escala industrial para a Cimed (Líder Nacional Farmacêutica).",
        "bullets": [
            "Stack: React 19, TypeScript, Vite, Framer Motion, Zustand e Supabase.",
            "Social Hub com rankings em tempo real, feed gamificado e missões de Drive-to-Store.",
            "Integração da IA 'Claud.ia' como assistente de saúde e produto; vídeo scraping e PWA nativo.",
            "Link do Projeto: https://cimedexperience.web.app/"
        ]
    },
    {
        "title": "Especialista em Automação & Data Engineering",
        "period": "2021 - Presente",
        "description": "Esteiras de extração de dados complexos e automação de processos críticos (LegalTech/GovTech).",
        "bullets": [
            "Scraper Imobiliário (2 estágios): HTTP + Playwright com Tor multi-porta e OCR customizado (PyTorch/CRNN).",
            "Certibot (LegalTech): Bot de emissão automática de certidões em 6+ portais (TRF3, TJSP, TST, Receita Federal).",
            "GIS Data Mining: Extração e normalização de polígonos de zoneamento a partir de documentos PDF/CAD.",
            "Automação de Certidões: Sistema Full-Stack com Supabase Edge Functions e InfoSimples API."
        ]
    }
]

SKILLS = {
    "Linguagens": "Javascript (React/Node), TypeScript, Angular, Python (Scraping/AI), Kotlin (Jetpack Compose), SQL.",
    "Cloud & Infra": "Supabase, Firebase, Docker, PostgreSQL, Linux Server, Git/GitHub, CI/CD.",
    "Data & Analytics": "Power BI, Geoprocessamento (GIS), Big Data, Web Scraping Anti-Bot.",
    "AI & Automation": "PyTorch (CRNN OCR), Playwright, ddddocr, Gemini API, Vertex AI, Computer Vision.",
    "Soft Skills": "Storytelling para Marketing Digital, Business English (Advanced), Gestão de Projetos Ágeis."
}

CERTIFICATIONS = [
    {"nome": "Angular", "org": "Loiane Training (28h)", "date": "Março 2026"},
    {"nome": "Storytelling para Marketing Digital", "org": "Santander / University of Chicago", "date": "Março 2026"},
    {"nome": "Business English, Part 1", "org": "Santander Open Academy", "date": "Fevereiro 2026"},
    {"nome": "Power BI", "org": "Santander Open Academy", "date": "Janeiro 2026"}
]

EDUCATION = [
    "Desenvolvimento de Sistemas - Especialização autônoma em Full-Stack e Engenharia de Dados de Alta Performance."
]

def create_docx(path):
    doc = Document()
    
    # Cabeçalho
    title = doc.add_heading(NAME, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{LOCATION} | {CONTACTS}\n").font.size = Pt(10)
    p.add_run(f"LinkedIn: {LINKEDIN} | GitHub: {GITHUB}").font.size = Pt(10)
    
    # Resumo
    doc.add_heading('Resumo Profissional', level=1)
    doc.add_paragraph(SUMMARY)
    
    # Experiência
    doc.add_heading('Experiência Profissional e Projetos', level=1)
    for exp in EXPERIENCES:
        p = doc.add_paragraph()
        run = p.add_run(exp['title'])
        run.bold = True
        run.font.size = Pt(12)
        p.add_run(f" ({exp['period']})")
        doc.add_paragraph(exp['description'])
        for bullet in exp['bullets']:
            doc.add_paragraph(bullet, style='List Bullet')
            
    # Expertise
    doc.add_heading('Expertise Técnica', level=1)
    for k, v in SKILLS.items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(v)
        
    # Certificações
    doc.add_heading('Certificações e Educação', level=1)
    for cert in CERTIFICACOES: # Usei o nome em PT no loop anterior, corrigindo para CERTIFICATIONS
        doc.add_paragraph(f"{cert['nome']} ({cert['org']}) - {cert['date']}", style='List Bullet')
    for edu in EDUCATION:
        doc.add_paragraph(edu, style='List Bullet')
        
    doc.save(path)

# Vou corrigir o loop de certificações no código final abaixo
def create_docx_final(path):
    doc = Document()
    title = doc.add_heading(NAME, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{LOCATION} | {CONTACTS}\n").font.size = Pt(10)
    p.add_run(f"LinkedIn: {LINKEDIN} | GitHub: {GITHUB}").font.size = Pt(10)
    
    doc.add_heading('Resumo Profissional', level=1)
    doc.add_paragraph(SUMMARY)
    
    doc.add_heading('Experiência Profissional e Projetos de Destaque', level=1)
    for exp in EXPERIENCES:
        p = doc.add_paragraph()
        run = p.add_run(exp['title'])
        run.bold = True
        run.font.size = Pt(12)
        p.add_run(f"  |  {exp['period']}")
        doc.add_paragraph(exp['description']).italic = True
        for bullet in exp['bullets']:
            doc.add_paragraph(bullet, style='List Bullet')
            
    doc.add_heading('Expertise Técnica (Hard Skills)', level=1)
    for k, v in SKILLS.items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(v)
        
    doc.add_heading('Certificações Recentes', level=1)
    for cert in CERTIFICATIONS:
        doc.add_paragraph(f"{cert['nome']} | {cert['org']} - {cert['date']}", style='List Bullet')
        
    doc.add_heading('Formação Acadêmica', level=1)
    for edu in EDUCATION:
        doc.add_paragraph(edu, style='List Bullet')
        
    doc.save(path)

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, NAME, 0, 1, 'C')
        self.set_font('Arial', '', 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, f"{LOCATION} | {CONTACTS}", 0, 1, 'C')
        self.cell(0, 5, f"LinkedIn: {LINKEDIN} | GitHub: {GITHUB}", 0, 1, 'C')
        self.ln(10)

    def section_title(self, title):
        self.set_font('Arial', 'B', 11)
        self.set_fill_color(240, 240, 240)
        self.set_text_color(0, 0, 0)
        self.cell(0, 8, title, 0, 1, 'L', 1)
        self.ln(3)

    def body_text(self, text, bullet=False, bold_prefix=""):
        self.set_font('Arial', '', 10)
        self.set_text_color(60, 60, 60)
        prefix = "- " if bullet else ""
        if bold_prefix:
            self.set_font('Arial', 'B', 10)
            self.write(6, f"{prefix}{bold_prefix}: ")
            self.set_font('Arial', '', 10)
            self.multi_cell(0, 6, text)
        else:
            self.multi_cell(0, 6, f"{prefix}{text}")
        self.ln(1)

def create_pdf_final(path):
    pdf = PDF()
    pdf.add_page()
    
    pdf.section_title('RESUMO PROFISSIONAL')
    pdf.body_text(SUMMARY)
    
    pdf.section_title('EXPERIÊNCIA PROFISSIONAL E PROJETOS')
    for exp in EXPERIENCES:
        pdf.set_font('Arial', 'B', 10)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 8, f"{exp['title']} ({exp['period']})", 0, 1)
        pdf.set_font('Arial', 'I', 9)
        pdf.cell(0, 5, exp['description'], 0, 1)
        for bullet in exp['bullets']:
            pdf.body_text(bullet, bullet=True)
        pdf.ln(2)
            
    pdf.section_title('EXPERTISE TÉCNICA')
    for k, v in SKILLS.items():
        pdf.body_text(v, bold_prefix=k)
        
    pdf.section_title('CERTIFICAÇÕES E FORMAÇÃO')
    for cert in CERTIFICATIONS:
        pdf.body_text(f"{cert['nome']} | {cert['org']} - {cert['date']}", bullet=True)
    
    for edu in EDUCATION:
        pdf.body_text(edu, bullet=True)
        
    pdf.output(path)

# Código de execução corrigido
if __name__ == "__main__":
    downloads_path = r"C:\Users\bruno\Downloads"
    output_docx = os.path.join(downloads_path, "Curriculo_Master_Bruno_Pereira.docx")
    output_pdf = os.path.join(downloads_path, "Curriculo_Master_Bruno_Pereira.pdf")
    
    create_docx_final(output_docx)
    create_pdf_final(output_pdf)
    print(f"Curriculo Master Gerado: {output_docx}")
    print(f"Curriculo Master Gerado: {output_pdf}")
