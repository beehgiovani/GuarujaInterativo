from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import os
import unicodedata

# --- DADOS DO CURRÍCULO ---
NAME = "Bruno Giovani Pereira"
LOCATION = "Guarujá, SP"
CONTACTS = "brunogp.corretor@gmail.com | (13) 99124-8146"
LINKEDIN = "https://www.linkedin.com/in/beehgiovani/"
GITHUB = "https://github.com/beehgiovani/"

PRIMARY_COLOR_RGB = (0, 74, 153)
PRIMARY_COLOR_HEX = RGBColor(0, 74, 153)

SUMMARY = (
    "Software Engineer Sênior com sólida expertise em arquiteturas de alta performance, especialização em Geotecnologia (GIS), "
    "Automação Avançada (OSINT/LegalTech) e Microsserviços. Experiência em sistemas de missão crítica, plataformas "
    "de escala industrial com processamento assíncrono e motores de extração de Big Data. Proficiência em ecossistemas Java/Kotlin (Spring Boot), "
    "otimização avançada de PostgreSQL e infraestrutura Cloud (AWS/Supabase). Certificações recentes pela University of Chicago e Loiane Training."
)

EXPERIENCES = [
    {
        "title": "Front-End Developer | PredictMed",
        "period": "Março 2026 - Presente",
        "description": "Projeto SaaS de saúde focado em automação, análise preditiva e infraestrutura serverless.",
        "bullets": [
            "Arquitetura moderna: React 19, TypeScript, Vite e Tailwind CSS para sistemas de alta complexidade.",
            "Backend as a Service: Integração com Supabase (PostgreSQL, Auth), Edge Functions (Deno/Typescript).",
            "Automação de fluxos complexos e integração de modelos de IA para análise preditiva de dados sensíveis.",
            "Foco em segurança de dados e UX de alto impacto com design responsivo e micro-animações.",
            "Deploy e CI/CD: Firebase Hosting e monitoramento de performance em tempo real."
        ]
    },
    {
        "title": "Senior Software Engineer (Kotlin/Java) | Soo Tech",
        "period": "Jan. 2024 - Presente",
        "description": "Consultoria sênior em ecossistema JVM para projetos de alta performance e escala.",
        "bullets": [
            "Desenvolvimento e otimização de módulos críticos em Java/Kotlin, focando em redução de latência e concorrência.",
            "Liderança técnica na implementação de APIs RESTful robustas utilizando Spring Boot e padrões de microsserviços.",
            "Garantia de qualidade via testes automatizados (JUnit, Mockito) e aplicação rigorosa de princípios SOLID e Clean Code.",
            "Colaboração em ambientes ágeis (Scrum) para entrega contínua de funcionalidades de alto valor de negócio.",
            "Desenvolvimento de integrações seguras e processamento de dados em ambientes de missão crítica."
        ]
    },
    {
        "title": "Founder & Lead Developer | GuaruGeo (Guarujá Geo Lab)",
        "period": "2023 - Presente",
        "description": "Plataforma GIS de inteligência imobiliária com backend PostgreSQL avançado.",
        "bullets": [
            "Arquitetura Full-Stack: React, Node.js e PostgreSQL com foco em performance e otimização de queries complexas.",
            "PostgreSQL Avançado: Implementação de Particionamento, Materialized Views, Triggers e índices espaciais (RBush/GIS).",
            "Desenvolvimento de motor OSINT para enriquecimento de dados via APIs seguras e processamento em lote (Batch).",
            "Dashboard de Analytics em tempo real com rastreamento de eventos e integração de auditoria de dados.",
            "Infraestrutura híbrida escalável e sistema automatizado de monitoramento de integridade de dados."
        ]
    },
    {
        "title": "Lead Developer | Cimed Experience",
        "period": "2024 - Março 2026",
        "description": "PWA de escala industrial para engajamento e gamificação (Líder Nacional Farmacêutica).",
        "bullets": [
            "Stack: React 19, TypeScript, Vite e infraestrutura cloud com foco em alta disponibilidade.",
            "Arquitetura baseada em eventos (Event-driven) para rankings em tempo real e feeds gamificados de alta concorrência.",
            "Integração de IA generativa (Gemini/Claude) para assistência de produto e automação de engajamento.",
            "Gestão de ambiente de larga escala com milhares de usuários ativos simultâneos e baixa latência."
        ]
    },
    {
        "title": "Especialista em Automação & Data Engineering",
        "period": "2021 - Presente",
        "description": "Esteiras de extração de dados (ETL) e automação de processos críticos (LegalTech/GovTech/Financial).",
        "bullets": [
            "Engineering de Dados: Extração e normalização de Big Data com segurança (Tor multi-porta, proxy rotation).",
            "LegalTech: Automação completa de certidões em 6+ portais governamentais com validação automática de PDFs.",
            "OCR Customizado: Treinamento de modelos deep learning (PyTorch/CRNN) para quebra de captchas e leitura de documentos.",
            "Integração de APIs de terceiros com tratamento de falhas (Circuit Breaker) e alta observabilidade."
        ]
    }
]

SKILLS = {
    "Linguagens": "Java (Spring Boot / Java 21), Kotlin (Jetpack Compose), Python (AI/Scraping), TypeScript, Javascript, SQL.",
    "Arquitetura & Dev": "Microsserviços, Arquitetura baseada em eventos, REST APIs, SOLID, Clean Code, JUnit, Mockito.",
    "Cloud & Bancos": "PostgreSQL (Otimização/Performance), AWS (S3/Lambda/DynamoDB), Supabase, Docker, GitHub Actions, CI/CD.",
    "Inteligência & Dados": "BI (Power BI), Geoprocessamento (GIS), Web Scraping Anti-Bot, PyTorch (Computer Vision).",
    "Soft Skills": "Storytelling para Negócios, Inglês Avançado, Gestão Ágil de Projetos de Software."
}

CERTIFICATIONS = [
    {"nome": "Angular", "org": "Loiane Training (28h)", "date": "26 de Março 2026"},
    {"nome": "Storytelling para Marketing Digital", "org": "Santander / University of Chicago", "date": "20 de Março 2026"},
    {"nome": "Business English, Part 1", "org": "Santander Open Academy", "date": "15 de Março 2026"},
    {"nome": "Power BI", "org": "Santander Open Academy", "date": "10 de Março 2026"}
]

EDUCATION = [
    "Bacharelado em Farmácia (Cursando) - Base para soluções HealthTech e Analytics Farmacêutico.",
    "Desenvolvimento de Sistemas - Especialização autônoma em Full-Stack e Engenharia de Dados de Alta Performance."
]

def remove_accents(input_str):
    if not isinstance(input_str, str): return input_str
    nfkd_form = unicodedata.normalize('NFKD', input_str)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)])

def create_docx_final(path):
    doc = Document()
    title = doc.add_heading(NAME, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs: run.font.color.rgb = PRIMARY_COLOR_HEX
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{LOCATION} | {CONTACTS}\n").font.size = Pt(10)
    p.add_run(f"LinkedIn: {LINKEDIN} | GitHub: {GITHUB}").font.size = Pt(10)
    
    sections = [
        ('Resumo Profissional', SUMMARY),
        ('Experiência Profissional e Projetos de Destaque', EXPERIENCES),
        ('Expertise Técnica (Hard Skills)', SKILLS),
        ('Certificações Recentes', CERTIFICATIONS),
        ('Formação Acadêmica', EDUCATION)
    ]
    for text, content in sections:
        h = doc.add_heading(text, level=1)
        for run in h.runs: run.font.color.rgb = PRIMARY_COLOR_HEX
        if text == 'Resumo Profissional': doc.add_paragraph(content)
        elif text == 'Experiência Profissional e Projetos de Destaque':
            for exp in content:
                p = doc.add_paragraph()
                run = p.add_run(exp['title'])
                run.bold, run.font.size, run.font.color.rgb = True, Pt(12), PRIMARY_COLOR_HEX
                p.add_run(f"  |  {exp['period']}")
                doc.add_paragraph(exp['description']).italic = True
                for bullet in exp['bullets']: doc.add_paragraph(bullet, style='List Bullet')
        elif text == 'Expertise Técnica (Hard Skills)':
            for k, v in content.items():
                p = doc.add_paragraph(); p.add_run(f"{k}: ").bold = True; p.add_run(v)
        elif text == 'Certificações Recentes':
            for cert in content: doc.add_paragraph(f"{cert['nome']} | {cert['org']} - {cert['date']}", style='List Bullet')
        elif text == 'Formação Acadêmica':
            for edu in content: doc.add_paragraph(edu, style='List Bullet')
    doc.save(path)

class PDF(FPDF):
    def header(self): pass
    def add_top_header(self):
        self.set_font('helvetica', 'B', 22)
        self.set_text_color(*PRIMARY_COLOR_RGB)
        self.cell(self.epw, 15, remove_accents(NAME), 0, 1, 'C')
        self.set_font('helvetica', '', 10)
        self.set_text_color(100, 100, 100)
        self.cell(self.epw, 7, f"{remove_accents(LOCATION)} | {remove_accents(CONTACTS)}", 0, 1, 'C')
        self.cell(self.epw, 7, f"LinkedIn: {LINKEDIN} | GitHub: {GITHUB}", 0, 1, 'C')
        self.ln(10)
    def section_title(self, title):
        self.set_font('helvetica', 'B', 12)
        self.set_fill_color(245, 248, 255) 
        self.set_text_color(*PRIMARY_COLOR_RGB)
        self.cell(self.epw, 10, f"  {remove_accents(title)}", 0, 1, 'L', 1)
        self.ln(2)
    def body_text(self, text, bullet=False, bold_prefix=""):
        self.set_font('helvetica', '', 10)
        self.set_text_color(50, 50, 50)
        prefix = "- " if bullet else ""
        clean_text = remove_accents(text)
        if bold_prefix:
            self.set_font('helvetica', 'B', 10)
            self.write(6, f"{prefix}{remove_accents(bold_prefix)}: ")
            self.set_font('helvetica', '', 10)
            self.multi_cell(self.epw, 6, clean_text)
        else:
            self.multi_cell(self.epw, 6, f"{prefix}{clean_text}")
        self.ln(1)

def create_pdf_final(path):
    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.add_top_header()
    
    sections = [
        ('RESUMO PROFISSIONAL', SUMMARY),
        ('EXPERIENCIA PROFISSIONAL E PROJETOS', EXPERIENCES),
        ('EXPERTISE TECNICA (HARD SKILLS)', SKILLS),
        ('CERTIFICACOES RECENTES', CERTIFICATIONS),
        ('FORMACAO ACADEMICA', EDUCATION)
    ]
    for title, content in sections:
        pdf.section_title(title)
        if title == 'RESUMO PROFISSIONAL': pdf.body_text(content)
        elif title == 'EXPERIENCIA PROFISSIONAL E PROJETOS':
            for exp in content:
                pdf.set_font('helvetica', 'B', 10); pdf.set_text_color(*PRIMARY_COLOR_RGB)
                pdf.cell(pdf.epw, 8, f"{remove_accents(exp['title'])}  |  {remove_accents(exp['period'])}", 0, 1)
                pdf.set_font('helvetica', 'I', 9); pdf.set_text_color(80, 80, 80)
                pdf.multi_cell(pdf.epw, 5, remove_accents(exp['description']))
                for b in exp['bullets']: pdf.body_text(b, bullet=True)
                pdf.ln(2)
        elif title == 'EXPERTISE TECNICA (HARD SKILLS)':
            for k, v in content.items(): pdf.body_text(v, bold_prefix=k)
        elif title == 'CERTIFICACOES RECENTES':
            for c in content: pdf.body_text(f"{c['nome']} | {c['org']} - {c['date']}", bullet=True)
        elif title == 'FORMACAO ACADEMICA':
            for e in content: pdf.body_text(e, bullet=True)

    pdf.output(path)

if __name__ == "__main__":
    downloads_path = r"C:\Users\bruno\Downloads"
    out_docx = os.path.join(downloads_path, "Curriculo_Premium_Bruno_Pereira.docx")
    out_pdf = os.path.join(downloads_path, "Curriculo_Premium_Bruno_Pereira.pdf")
    create_docx_final(out_docx)
    create_pdf_final(out_pdf)
    print(f"Gerado: {out_docx}\nGerado: {out_pdf}")
