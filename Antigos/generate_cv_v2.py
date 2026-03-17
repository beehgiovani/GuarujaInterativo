from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import os

# CONFIGURAÇÕES DE CONTATO (Conforme solicitado)
NOME = "Bruno Giovani Pereira"
DATA_NASC = "25/05/1993"
LOCAL = "Guarujá, SP"
EMAIL = "brunogp.corretor@gmail.com"
WHATSAPP = "(13) 99123-4567"  # Substitua pelo seu número real se preferir
LINKEDIN = "linkedin.com/in/bruno-giovani-pereira" # Substitua pelo seu link real

def create_docx(filename):
    doc = Document()
    
    # Estilo de Título
    title = doc.add_heading(NOME, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações de Contato
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{LOCAL} | {DATA_NASC}\n{EMAIL} | {WHATSAPP}\n{LINKEDIN}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    # RESUMO
    doc.add_heading('Resumo Profissional', level=1)
    doc.add_paragraph(
        "Desenvolvedor Full-Stack, Especialista em Automação (OSINT) e Consultor Imobiliário. "
        "Perfil híbrido único que combina domínio técnico avançado (IA, Web Scraping, Geoprocessamento) "
        "com profunda visão de mercado imobiliário e farmacêutico. Especialista em transformar dados complexos "
        "em ferramentas de alta performance para tomada de decisão e geração de receita."
    )
    
    # HARD SKILLS
    doc.add_heading('Expertise Técnica (Hard Skills)', level=1)
    skills = [
        "Desenvolvimento: JavaScript (Node.js/ES6+), Python, Kotlin/Java (Android Native).",
        "Mobile & Desktop: Jetpack Compose, Electron, Capacitor, Flutter (Básico).",
        "Inteligência Artificial & Dados: Reinforcement Learning (DQN), Séries Temporais, OCR (ddddocr), Integração LLM (Gemini/OpenAI).",
        "Geoprocessamento: Google Maps Platform (Advanced), Google Earth Engine, Leaflet, ArcGIS/OSM.",
        "Automação & OSINT: Selenium, Playwright, Stealth Scraping, Bypass de Captcha, Agentes Autônomos.",
        "Infraestrutura: Firebase, Supabase (PostgreSQL), Docker, Git/GitHub, CI/CD."
    ]
    for skill in skills:
        doc.add_paragraph(skill, style='List Bullet')
        
    # PROJETOS DE DESTAQUE
    doc.add_heading('Ecossistema de Projetos & Inovação', level=1)
    
    # 1. Guarujá Geo Lab
    doc.add_heading('1. Guarujá Geo Lab / GeoMap (GovTech/PropTech)', level=2)
    doc.add_paragraph(
        "Plataforma líder em inteligência territorial para o litoral paulista. "
        "Integra dados de zoneamento, mapas 3D e CRM imobiliário. Automatiza a prospecção de áreas, "
        "identificando proprietários e débitos fiscais em segundos através de motores OSINT proprietários."
    )
    
    # 2. MarketPoster 3.0
    doc.add_heading('2. MarketPoster 3.0 (Varejo & Design)', level=2)
    doc.add_paragraph(
        "Aplicativo Android profissional (Jetpack Compose) para gestão de sinalização de varejo. "
        "Inclui processamento de imagem via ML Kit (Background Removal), sincronização offline redundante "
        "e suporte a impressão térmica Bluetooth."
    )

    # 3. Slot Automator Ultra (IA & RL)
    doc.add_heading('3. IA & Reinforcement Learning Agent', level=2)
    doc.add_paragraph(
        "Experimento científico avançado utilizando Deep Q-Networks (DQN) e Análise de Séries Temporais "
        "para tomada de decisão em ambientes de recompensa variável. Inclui módulos de visão computacional "
        "e gestão de risco matemática (Critério de Kelly)."
    )

    # 4. Bot de Certidões Automatizadas
    doc.add_heading('4. Automação de Processos Jurídicos (LegalTech)', level=2)
    doc.add_paragraph(
        "Desenvolvimento de robôs (Playwright) para emissão automática de certidões em portais do governo "
        "(TJSP, TRT, CJF, RFB). Implementação de sistema de OCR gratuito offline para resolução de desafios visuais."
    )

    # 5. PharmaFlow
    doc.add_heading('5. PharmaFlow (HealthTech)', level=2)
    doc.add_paragraph(
        "Sistema em desenvolvimento voltado para otimização de fluxos farmacêuticos e gestão de estoque inteligente, "
        "unindo o conhecimento acadêmico em Farmácia com a expertise em desenvolvimento mobile."
    )
    
    # EXPERIÊNCIA
    doc.add_heading('Trajetória Profissional', level=1)
    doc.add_paragraph("Founder & Lead Developer | Guarujá Geo Lab (2023 – Presente)", style='List Bullet')
    doc.add_paragraph("Especialista em Inteligência Imobiliária | Litoral de SP", style='List Bullet')
    doc.add_paragraph("Desenvolvedor Independente de Automações & Bots OSINT", style='List Bullet')
    
    # EDUCAÇÃO
    doc.add_heading('Educação', level=1)
    doc.add_paragraph("Bacharelado em Farmácia | [Nome da Instituição] (Cursando)", style='List Bullet')
    doc.add_paragraph("Desenvolvimento de Sistemas | Autodidata com Especializações Online", style='List Bullet')
    
    doc.save(filename)

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 16)
        self.cell(0, 10, NOME, 0, 1, 'C')
        self.set_font('Arial', '', 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, f'{LOCAL} | {DATA_NASC}', 0, 1, 'C')
        self.cell(0, 5, f'{EMAIL} | {WHATSAPP}', 0, 1, 'C')
        self.cell(0, 5, LINKEDIN, 0, 1, 'C')
        self.ln(10)

    def chapter_title(self, title):
        self.set_font('Arial', 'B', 11)
        self.set_fill_color(240, 240, 240)
        self.set_text_color(0, 0, 0)
        self.cell(0, 8, title, 0, 1, 'L', 1)
        self.ln(3)

    def chapter_body(self, body, bold_title="", bullet=False):
        self.set_font('Arial', '', 10)
        self.set_text_color(60, 60, 60)
        prefix = "- " if bullet else ""
        if bold_title:
            self.set_font('Arial', 'B', 10)
            self.write(6, f"{prefix}{bold_title}: ")
            self.set_font('Arial', '', 10)
            self.multi_cell(0, 6, body)
        else:
            self.multi_cell(0, 6, f"{prefix}{body}")
        self.ln(2)

def create_pdf(filename):
    pdf = PDF()
    pdf.add_page()
    
    pdf.chapter_title('RESUMO PROFISSIONAL')
    pdf.chapter_body(
        "Desenvolvedor Full-Stack, Especialista em Automação (OSINT) e Consultor Imobiliário. "
        "Perfil híbrido único que combina domínio técnico avançado (IA, Web Scraping, Geoprocessamento) "
        "com profunda visão de mercado imobiliário e farmacêutico."
    )
    
    pdf.chapter_title('EXPERTISE TÉCNICA')
    skills = [
        "JS (Node/React), Python, Kotlin (Android Native), Jetpack Compose.",
        "Reinforcement Learning, Séries Temporais, OCR, IA Generativa (Gemini).",
        "Google Maps API, GIS, Automação OSINT (Playwright/Selenium).",
        "Firebase, Supabase (PostgreSQL), Cloud Hosting."
    ]
    for skill in skills:
        pdf.chapter_body(skill, bullet=True)
        
    pdf.chapter_title('PROJETOS E INOVAÇÃO')
    pdf.chapter_body("Plataforma GovTech de geoprocessamento e inteligência territorial.", bold_title="Guarujá Geo Lab")
    pdf.chapter_body("App Android nativo para automação de design e impressão no varejo.", bold_title="MarketPoster 3.0")
    pdf.chapter_body("Agente de IA baseado em Deep Q-Networks para ambientes variáveis.", bold_title="Slot IA Manager")
    pdf.chapter_body("Automação de certidões jurídicas e resolução de desafios visuais (OCR).", bold_title="Legal Bot")
    pdf.chapter_body("Sistema móvel para gestão de fluxos em farmácias e drogarias.", bold_title="PharmaFlow")
    
    pdf.chapter_title('FORMAÇÃO E TRAJETÓRIA')
    pdf.chapter_body("Founder / Lead Developer no Guarujá Geo Lab.", bullet=True)
    pdf.chapter_body("Superior em Farmácia (Cursando).", bullet=True)
    pdf.chapter_body("Desenvolvedor Autodidata focado em soluções Enterprise.", bullet=True)
    
    pdf.output(filename)

if __name__ == "__main__":
    create_docx("Curriculo_Bruno_Pereira_V2.docx")
    create_pdf("Curriculo_Bruno_Pereira_V2.pdf")
    print("CV V2 atualizado com sucesso!")
.gitignore
