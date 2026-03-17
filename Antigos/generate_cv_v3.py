from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import os

# CONFIGURAÇÕES DE CONTATO
NOME = "Bruno Giovani Pereira"
DATA_NASC = "25/05/1993"
LOCAL = "Guarujá, SP"
EMAIL = "brunogp.corretor@gmail.com"
WHATSAPP = "(13) 99123-4567" 
LINKEDIN = "linkedin.com/in/bruno-giovani-pereira"

def create_docx(filename):
    doc = Document()
    
    # Estilo de Título
    title = doc.add_heading(NOME, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informações de Contato
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{LOCAL} | {DATA_NASC}\n{EMAIL} | {WHATSAPP}\n{LINKEDIN}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    # RESUMO
    doc.add_heading('Resumo Profissional', level=1)
    doc.add_paragraph(
        "Desenvolvedor Full-Stack e Especialista em Inteligência de Dados (OSINT). "
        "Destaque na criação de soluções de alta complexidade que unem tecnologia avançada e visão estratégica de negócios. "
        "Experiência sólida no desenvolvimento de sistemas GIS (Geoprocessamento), automação de processos governamentais/judiciais, "
        "aplicativos móveis nativos e agentes de Inteligência Artificial baseados em Aprendizado por Reforço."
    )
    
    # HARD SKILLS
    doc.add_heading('Expertise Técnica (Hard Skills)', level=1)
    skills = [
        "Linguagens & Frameworks: JavaScript (Node.js/ES6+), Python, Kotlin (Android Jetpack Compose), Java.",
        "Mobile & Desktop: Electron (Desktop App), Capacitor, Mobile Native Android.",
        "Inteligência Artificial: Reinforcement Learning (DQN), Séries Temporais, Visão Computacional, OCR (Offline/Neural).",
        "Geoprocessamento: Google Maps Platform (Advanced), Google Earth Engine, OpenStreetMap (OSM) Integrations.",
        "Automação & OSINT: Selenium, Playwright, Stealth Scraping, Bypass de Captchas, Multi-Stage Scrapers.",
        "Arquitetura de Dados: Firebase (Firestore), Supabase (PostgreSQL), REST APIs, Real-time Sync."
    ]
    for skill in skills:
        doc.add_paragraph(skill, style='List Bullet')
        
    # PROJETOS DE DESTAQUE
    doc.add_heading('Projetos e Soluções Tecnológicas', level=1)
    
    # 1. Guarujá Geo Lab
    doc.add_heading('1. Plataforma de Inteligência Territorial & GIS (Guarujá Geo Lab)', level=2)
    doc.add_paragraph(
        "Ecossistema de geoprocessamento escalável que integra zoneamento municipal, visualização 3D e CRM. "
        "Atua como um hub central de dados imobiliários, automatizando a prospecção ativa através do cruzamento de bases "
        "fiscais e territoriais em tempo real."
    )
    
    # 2. Solução Enterprise de Design & PDV Mobile
    doc.add_heading('2. Ecossistema Mobile para Gestão de Varejo (PDV Digital)', level=2)
    doc.add_paragraph(
        "Desenvolvimento nativo Android (Jetpack Compose) focado na automação de comunicação visual em pontos de venda. "
        "Integra processamento de imagem com ML Kit para remoção inteligente de fundo (Background Removal) e comunicação "
        "direta via Bluetooth com equipamentos de impressão térmica."
    )

    # 3. Agente de Tomada de Decisão em Tempo Real (DQN Agent)
    doc.add_heading('3. Autonomous Decision Agent (Reinforcement Learning)', level=2)
    doc.add_paragraph(
        "Implementação de uma Super-IA baseada em Aprendizado por Reforço Profundo (DQN) aplicada a análise de dados "
        "dinâmicos. Utiliza ciência de dados avançada e análise de séries temporais para predição de ciclos e gestão "
        "matemática de risco baseada no Critério de Kelly."
    )

    # 4. Sistema Automatizado de OSINT Governamental & Documental
    doc.add_heading('4. Legal & Gov OSINT Bot (Automação Documental)', level=2)
    doc.add_paragraph(
        "Robô de automação inteligente (Playwright/Python) para recuperação em massa de certidões e documentos em portais "
        "judiciais e federais. Utiliza motores de OCR neurais para resolução automática de desafios visuais, "
        "eliminando a necessidade de intervenção humana em processos burocráticos."
    )

    # 5. Ecossistema de Gestão HealthTech
    doc.add_heading('5. Plataforma de Otimização de Fluxos em Saúde (HealthTech)', level=2)
    doc.add_paragraph(
        "Sistema dedicado à gestão inteligente de estoques e fluxos de atendimento farmacêutico, unindo análise "
        "clínica e tecnologia móvel para redução de gargalos operacionais no setor de saúde."
    )
    
    # TRAJETÓRIA
    doc.add_heading('Trajetória e Atuação Freelance', level=1)
    doc.add_paragraph("Founder & Lead Developer | Guarujá Geo Lab (2023 – Presente)", style='List Bullet')
    doc.add_paragraph("Desenvolvedor Freelance de Automações OSINT e Soluções Mobile", style='List Bullet')
    doc.add_paragraph("Consultor de Inteligência Territorial para Incorporação Imobiliária", style='List Bullet')
    
    # EDUCAÇÃO
    doc.add_heading('Educação', level=1)
    doc.add_paragraph("Superior em Farmácia | [Instituição] (Cursando)", style='List Bullet')
    doc.add_paragraph("Desenvolvimento de Software Full-Stack (Formação Autodidata Avançada)", style='List Bullet')
    
    doc.save(filename)

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 16)
        self.cell(0, 10, NOME, 0, 1, 'C')
        self.set_font('Arial', '', 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, f'{LOCAL} | {DATA_NASC}', 0, 1, 'C')
        self.cell(0, 5, f'{EMAIL} | {WHATSAPP}', 0, 1, 'C')
        self.ln(10)

    def chapter_title(self, title):
        self.set_font('Arial', 'B', 11)
        self.set_fill_color(240, 240, 240)
        self.set_text_color(0, 0, 0)
        self.cell(0, 8, title, 0, 1, 'L', 1)
        self.ln(3)

    def chapter_body(self, body, bold_title="", bullet=False):
        self.set_font('Arial', '', 10)
        self.set_text_color(60, 60, 60)
        prefix = "- " if bullet else ""
        if bold_title:
            self.set_font('Arial', 'B', 10)
            self.write(6, f"{prefix}{bold_title}: ")
            self.set_font('Arial', '', 10)
            self.multi_cell(0, 6, body)
        else:
            self.multi_cell(0, 6, f"{prefix}{body}")
        self.ln(2)

def create_pdf(filename):
    pdf = PDF()
    pdf.add_page()
    
    pdf.chapter_title('RESUMO PROFISSIONAL')
    pdf.chapter_body(
        "Desenvolvedor Full-Stack e Especialista em Inteligência de Dados (OSINT). "
        "Criação de soluções de alta complexidade em Geoprocessamento (GIS), Automação Documental e Inteligência Artificial."
    )
    
    pdf.chapter_title('EXPERTISE TÉCNICA (HARD SKILLS)')
    pdf.chapter_body("Node.js, Python, Kotlin (Android Jetpack Compose), Java.", bullet=True)
    pdf.chapter_body("Deep Q-Networks (RL), Séries Temporais, OCR Neural, ML Kit.", bullet=True)
    pdf.chapter_body("Google Maps Advanced, Automations (Playwright/Stealth), GIS.", bullet=True)
    pdf.chapter_body("Firebase, Supabase (PostgreSQL), Real-time Synchronization.", bullet=True)
        
    pdf.chapter_title('PRINCIPAIS PROJETOS')
    pdf.chapter_body("Hub de geoprocessamento e inteligência territorial com CRM.", bold_title="Guarujá Geo Lab")
    pdf.chapter_body("Solução mobile nativa para design PDV e automação de impressão.", bold_title="Retail Mobile Enterprise")
    pdf.chapter_body("Agente autônomo baseado em Reinforcement Learning (DQN).", bold_title="Decision Engine AI")
    pdf.chapter_body("Recuperação automatizada de dados judiciais com OCR offline.", bold_title="Legal OSINT Bot")
    pdf.chapter_body("Ecossistema mobile para gestão de fluxos de saúde (HealthTech).", bold_title="Health Analytics App")
    
    pdf.chapter_title('FORMAÇÃO E ATUAÇÃO')
    pdf.chapter_body("Founder & Lead Developer no Guarujá Geo Lab.", bullet=True)
    pdf.chapter_body("Superior em Farmácia (Cursando).", bullet=True)
    pdf.chapter_body("Foco em soluções corporativas de alta performance.", bullet=True)
    
    pdf.output(filename)

if __name__ == "__main__":
    create_docx("Curriculo_Bruno_Pereira_Final.docx")
    create_pdf("Curriculo_Bruno_Pereira_Final.pdf")
    print("Versão Final Profissional gerada!")
.gitignore
